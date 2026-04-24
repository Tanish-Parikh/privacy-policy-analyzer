import copy
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


INPUT_PATH = Path(r"C:\Users\Tanish Parikh\Downloads\Report.docx")
OUTPUT_PATH = Path(r"C:\Users\Tanish Parikh\Downloads\Report_formatted.docx")


def set_run_font(run, size=12, bold=False):
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold


def set_paragraph_spacing(paragraph):
    paragraph.paragraph_format.line_spacing = 1.5


def find_project_title(doc):
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if "PRIVACY POLICY ANALYZER" in text.upper():
            return "PRIVACY POLICY ANALYZER"
    return "Project Title"


def find_academic_year(doc):
    for paragraph in doc.paragraphs:
        match = re.search(r"\(\s*A\.Y\.\s*[^)]+\)", paragraph.text, flags=re.I)
        if match:
            return match.group(0)
    return "Academic Year"


def insert_section_break_before_introduction(doc):
    target = None
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if index > 50 and re.match(r"^2\.\s+Introduction\s*$", text, flags=re.I):
            target = paragraph
            break

    if target is None:
        return

    body = doc._body._body
    body_section_properties = body.sectPr

    new_paragraph_xml = OxmlElement("w:p")
    paragraph_properties = OxmlElement("w:pPr")
    new_section_properties = copy.deepcopy(body_section_properties)

    section_type = new_section_properties.find(qn("w:type"))
    if section_type is None:
        section_type = OxmlElement("w:type")
        new_section_properties.append(section_type)
    section_type.set(qn("w:val"), "nextPage")

    paragraph_properties.append(new_section_properties)
    new_paragraph_xml.append(paragraph_properties)
    target._p.addprevious(new_paragraph_xml)


def add_page_number_field(paragraph):
    for run in list(paragraph.runs):
        paragraph._p.remove(run._element)

    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    set_run_font(run, size=12, bold=False)

    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    # Force Arabic numerals in footer display even if section numbering is Roman.
    instr_text.text = "PAGE \\* Arabic"

    field_separator = OxmlElement("w:fldChar")
    field_separator.set(qn("w:fldCharType"), "separate")

    field_text = OxmlElement("w:t")
    field_text.text = "1"

    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")

    run._r.append(field_begin)
    run._r.append(instr_text)
    run._r.append(field_separator)
    run._r.append(field_text)
    run._r.append(field_end)


def configure_sections(doc, project_title, academic_year):
    sections = doc.sections
    if len(sections) >= 2:
        sections[1].start_type = WD_SECTION_START.NEW_PAGE

    for index, section in enumerate(sections):
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

        section_properties = section._sectPr
        page_number_type = section_properties.find(qn("w:pgNumType"))
        if page_number_type is None:
            page_number_type = OxmlElement("w:pgNumType")
            section_properties.append(page_number_type)

        if index == 0 and len(sections) >= 2:
            page_number_type.set(qn("w:fmt"), "lowerRoman")
            page_number_type.set(qn("w:start"), "1")
        else:
            page_number_type.set(qn("w:fmt"), "decimal")
            if len(sections) >= 2:
                page_number_type.set(qn("w:start"), "1")

        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False

        header_paragraph = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
        for run in list(header_paragraph.runs):
            header_paragraph._p.remove(run._element)
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_paragraph.paragraph_format.tab_stops.clear_all()
        tab_position = section.page_width - section.left_margin - section.right_margin
        header_paragraph.paragraph_format.tab_stops.add_tab_stop(tab_position, WD_TAB_ALIGNMENT.RIGHT)

        left_run = header_paragraph.add_run(project_title)
        tab_run = header_paragraph.add_run("\t")
        right_run = header_paragraph.add_run(academic_year)
        for run in (left_run, tab_run, right_run):
            set_run_font(run, size=12, bold=False)

        footer_paragraph = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        add_page_number_field(footer_paragraph)


def is_figure_caption(text):
    return bool(re.match(r"^\[?\s*Figure\s+([0-9]+|[IVXLC]+)\b", text, flags=re.I))


def is_table_caption(text):
    return bool(re.match(r"^\[?\s*Table\s+([0-9]+|[IVXLC]+)\b", text, flags=re.I))


def insert_paragraph_before(paragraph, text):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = paragraph.style
    new_para.add_run(text)
    return new_para


def ensure_dot_leader_with_page(line_text, page_label):
    clean = line_text.rstrip()
    match = re.match(r"^(.*?\.{5,})\s*([ivxlcdmIVXLCDM]+|\d+)?\s*$", clean)
    if match:
        return f"{match.group(1)} {page_label}"

    title = re.sub(r"\s+$", "", clean)
    dots = "." * max(20, 98 - len(title))
    return f"{title}{dots} {page_label}"


def apply_text_styles(doc):
    toc_custom_entries = [
        ("List of Figures", "i"),
        ("List of Tables", "ii"),
        ("Abbreviations", "iii"),
    ]

    toc_page_map = {
        "1.  Abstract": "iv",
        "2.  Introduction": "1",
        "2.1  Background and Motivation": "2",
        "2.2  Objectives": "3",
        "3.  Related Work and Literature Survey": "4",
        "3.1  Readability and Longitudinal Trends": "5",
        "3.2  Automated Classification and NLP Approaches": "6",
        "3.3  User Awareness and Visualisation": "7",
        "3.4  Annotation Variability and Explainability": "8",
        "3.5  Literature Survey Table": "9",
        "3.6  Research Gap": "10",
        "4.  Materials and Methods": "11",
        "4.1  System Architecture Overview": "12",
        "4.2  Data Collection and Usage": "13",
        "4.3  Algorithm: Risk Detection Engine": "14",
        "4.4  Algorithm: Readability Scoring (Flesch Reading Ease)": "15",
        "4.5  Algorithm: Privacy Risk Scoring": "16",
        "4.6  AI Clause Simplification (Gemini 2.5 Flash)": "17",
        "4.7  Proactive Badge Injection": "18",
        "4.8  Technology Stack": "19",
        "5.  Results (Screenshots)": "20",
        "6.  Conclusion": "21",
        "7.  Future Scope": "22",
        "8.  References": "23",
    }

    style_names = {style.name for style in doc.styles}
    if "Normal" in style_names:
        normal = doc.styles["Normal"]
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(12)

    toc_start = None
    toc_end = None
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text.upper() == "TABLE OF CONTENTS":
            toc_start = i
            continue
        if toc_start is not None and re.match(r"^1\.\s+Abstract\s*$", text, flags=re.I):
            toc_end = i
            break

    if toc_start is not None and toc_end is not None:
        first_entry_paragraph = doc.paragraphs[toc_end]
        toc_existing_text = [p.text.strip() for p in doc.paragraphs[toc_start:toc_end]]
        for entry_title, entry_page in reversed(toc_custom_entries):
            if entry_title not in toc_existing_text:
                insert_paragraph_before(
                    first_entry_paragraph,
                    ensure_dot_leader_with_page(entry_title, entry_page),
                )

        # Recompute TOC bounds after potential insertions.
        toc_start = None
        toc_end = None
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text.upper() == "TABLE OF CONTENTS":
                toc_start = i
                continue
            if toc_start is not None and re.match(r"^1\.\s+Abstract\b", text, flags=re.I):
                toc_end = i
                break

    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        set_paragraph_spacing(paragraph)
        if not paragraph.runs:
            continue

        in_toc_block = toc_start is not None and toc_end is not None and toc_start <= i < toc_end
        if in_toc_block:
            # Remove spacer lines in TOC that can appear as stray dot-leader rows.
            if text == "":
                paragraph.text = ""
                continue

            normalized = re.sub(r"\s*[ivxlcdmIVXLCDM\d]+\s*$", "", text).rstrip()
            normalized = normalized.lstrip("\t")

            page_label = None
            if normalized in toc_page_map:
                page_label = toc_page_map[normalized]
            else:
                for entry_title, entry_page in toc_custom_entries:
                    if normalized.startswith(entry_title):
                        page_label = entry_page
                        break

            if page_label is not None:
                indentation = "\t" if paragraph.text.startswith("\t") else ""
                paragraph.text = indentation + ensure_dot_leader_with_page(normalized, page_label)

            toc_text = paragraph.text.strip()
            is_toc_title = toc_text.upper() == "TABLE OF CONTENTS"

            for run in paragraph.runs:
                # Keep TOC typography consistent and prevent TOC rows from being styled as chapter headings.
                set_run_font(run, size=12, bold=is_toc_title or bool(run.font.bold))
            continue

        toc_entry = bool(re.search(r"\.{5,}", text))
        figure_caption = is_figure_caption(text)
        table_caption = is_table_caption(text)

        chapter = bool(re.match(r"^\d+\.\s+", text)) and not toc_entry
        section = bool(re.match(r"^\d+\.\d+\s+", text)) and not toc_entry
        subsection = bool(re.match(r"^\d+\.\d+\.\d+\s+", text)) and not toc_entry

        if (section or subsection) and len(text) > 120:
            split = re.match(r"^(\d+\.\d+(?:\.\d+)?\s+[^.]+\.)(\s+.*)$", text)
            if split:
                for run in list(paragraph.runs):
                    paragraph._p.remove(run._element)

                heading_run = paragraph.add_run(split.group(1))
                heading_size = 14 if section and not subsection else 12
                set_run_font(heading_run, size=heading_size, bold=True)

                body_run = paragraph.add_run(split.group(2))
                set_run_font(body_run, size=12, bold=False)
                continue

        if figure_caption or table_caption:
            for run in paragraph.runs:
                set_run_font(run, size=10, bold=False)
        elif chapter and not section and not subsection and len(text) < 100:
            for run in paragraph.runs:
                set_run_font(run, size=18, bold=True)
        elif section and not subsection and len(text) < 100:
            for run in paragraph.runs:
                set_run_font(run, size=14, bold=True)
        elif subsection and len(text) < 100:
            for run in paragraph.runs:
                set_run_font(run, size=12, bold=True)
        else:
            for run in paragraph.runs:
                set_run_font(run, size=12, bold=bool(run.font.bold))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    set_paragraph_spacing(paragraph)
                    for run in paragraph.runs:
                        set_run_font(run, size=12, bold=bool(run.font.bold))


def main():
    document = Document(str(INPUT_PATH))
    project_title = find_project_title(document)
    academic_year = find_academic_year(document)

    insert_section_break_before_introduction(document)
    configure_sections(document, project_title, academic_year)
    apply_text_styles(document)

    document.save(str(OUTPUT_PATH))
    print(f"Saved formatted document: {OUTPUT_PATH}")
    print(f"Sections: {len(document.sections)}")


if __name__ == "__main__":
    main()
